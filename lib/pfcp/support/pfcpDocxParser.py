#!/usr/bin/env python3

# Please save the doc as docx before delete useless table.
# Check all of table are complete. There are problems if rows are not align in table.
# All tables are saved in variable "tables" using structure "list".
# Rows for each table use structure "dict" and save in variable "tables[index]".

import os, re
from docx import Document

ieNameMapping = {
    'PFD': 'PFDContext',
    'PDRID': 'PacketDetectionRuleID',
    'SxSRRspFlags': 'PFCPSRRspFlags'
}


def snack2CamalCast(name):
    return ''.join(x.title() for x in name.lower().split('_'))

def large2smallCamalCast(name):
    return name[0].lower() + name[1:]

def formatString(inputStr : str) -> str :
    inputStr = re.sub(r"['/\"]", '', inputStr)
    outputStrList = [s[0].upper() + s[1:] for s in re.sub(r'[-() ]+', ' ', inputStr).split()]
    return ''.join(outputStrList)

def snackCase(inputStr: str) -> str:
    s1 = re.sub('(.)([A-Z][a-z]+)', r'\1_\2', inputStr) 
    return re.sub('([a-z0-9])([A-Z])', r'\1_\2', s1).upper().replace(" ", "")

class FileOutput() :
    def __init__(self, fileName : str) :
        self.fd = open(fileName, 'w')
        self.tab, self.tabstop = 0, 4

    def indent(self, num : int) :
        self.tab += num
        self.tab = 0 if self.tab < 0 else self.tab

    def indented(self, contentList : list) :
        self.indent(+1)
        for content in contentList :
            self.fprint(content)
        self.indent(-1)

    def fprint(self, content : str) :
        print(' ' * self.tabstop * self.tab, content, sep='', file=self.fd)

class TableParser() :
    def __init__(self, fileName : str) :
        self.document = Document(fileName)
        self.tables = []
        self.parse()

    def parse(self) :
        for idx, table in enumerate(self.document.tables) :
            gotTitle, titlePointer = 0, None
            for row in table.rows :
                try :
                    if 'Information elements'.lower() in [cell.text.lower() for cell in row.cells] :
                        if gotTitle == 0 :
                            self.tables.append(list())
                        titlePointer, gotTitle = row, 1
                    elif gotTitle == 1 :
                        content, isNote = dict(), 0
                        for title, context in zip(self.yieldTitleFromDocx(titlePointer), row.cells) :
                            if context._tc.right - context._tc.left >= 8 :
                                isNote = 1
                                break
                            content[title] = context.text
                        if isNote == 0 :
                            self.tables[-1].append(content)
                except :
                    print(f'[Error] The {idx} table is dirty')
                    break

    def yieldTitleFromDocx(self, tableRowPtr) :
        for cell in tableRowPtr.cells :
            yield cell.text

    def printTableByIndex(self, idxOfTable) :
        try :
            for content in self.tables[idxOfTable] :
                print(content)
        except :
            print('[Warning] Index out of bound')


if __name__ == '__main__' :

    doc29244_812_1 = TableParser('29244-f30-ch8.1.2-1.docx')
    ie_type_value = dict()
    for row in doc29244_812_1.tables[0][:-1]:
        ieName = formatString(row['Information elements'])
        if ieName == 'UpdateBARSessionModificationRequest':
            ieName = 'UpdateBARPFCPSessionModificationRequest'
        ieVal = row['IE Type value\n(Decimal)']
        if ie_type_value.get(ieName) == None:
            ie_type_value[ieName] = int(ieVal)
        else :
            print(f'[Warning] {ieName} is duplicate')

    specialCase = set()
    specialCase.update(['UpdateBAR', 'UsageReport'])

    # There have 67 table in chapter 7, but the first one will not be used
    docxChapter7Name = '29244-f30-ch7-fixed-table.docx'
    doc29244_7_para = Document(docxChapter7Name)
#    tableName = re.compile(r'Table 7.*: (Information Elements in [an ]{0,3})?(.+(?= IE within ))?(.+)')
    tableName = re.compile(r'Table 7.*: (Information Elements in [an ]{0,3}|(.+)( IE within ))?(.+)')
    chapter7TitleList = []
    for line in doc29244_7_para.paragraphs :
        afterMatch = tableName.match(line.text)
        if afterMatch :
            ieName = afterMatch.group(2) if afterMatch.group(2) else afterMatch.group(4)
            if formatString(ieName) in specialCase :
                ieName += afterMatch.group(4)
            chapter7TitleList.append(ieName)
            # print(afterMatch.group(2)) if afterMatch.group(2) else print(afterMatch.group(3))

    doc29244_7 = TableParser(docxChapter7Name)
    chapter7UsedIESet = set()
    for tableName in chapter7TitleList[1:] :
        tableName = formatString(tableName)
        ieIn = re.compile("^.*IEIn.*")
        if tableName == "UpdateBARIEInPFCPSessionReportResponse":
            tableName = "UpdateBARPFCPSessionReportResponse"
        elif ieIn.match(tableName):
            #print("============", tableName, tableName[:tableName.find("IEIn")])
            tableName = tableName[:tableName.find("IEIn")]
        elif tableName == 'RemoveQERIEPFCPSessionModificationRequest':
            tableName = tableName[:tableName.find("IE")]
        chapter7UsedIESet.add(tableName)

    PFCPMessageHeaderFd = FileOutput('pfcp_message.h')
    PFCPMessageHeaderFd.fprint('''#ifndef __PFCP_MESSAGE_H__
#define __PFCP_MESSAGE_H__

#include <stdint.h>

#include "utlt_debug.h"
#include "utlt_lib.h"
#include "utlt_buff.h"

#ifdef __cplusplus
extern "C" {
#endif /* __cplusplus */

typedef struct _TlvOctet {
    unsigned long presence;
    uint16_t type;
    uint16_t len;
    void *value;
} __attribute__((packed)) TlvOctet;

typedef struct _IeDescription {
    uint16_t msgType;
    uint16_t msgLen; // msg struct size
    _Bool isTlvObj;
    int numToParse;
    int next[35];
} __attribute__((packed)) IeDescription;

/* 5.1 General format */
#define PFCP_HEADER_LEN     16
#define PFCP_SEID_LEN       8
typedef struct _PfcpHeader {
    union {
        struct {
            ENDIAN4(uint8_t version:3;,
                    uint8_t spare0:3;,
                    uint8_t mp:1;,
                    uint8_t seidP:1;)
        };
        uint8_t flags;
    };
    uint8_t type;
    uint16_t length;
    union {
        struct {
            uint64_t seid;
#define PfcpTransactionId2Sqn(__transactionId) htonl(((__transactionId) << 8))
#define PfcpSqn2TransactionId(__sqn) (ntohl(__sqn) >> 8)
            uint32_t sqn;
        };
        uint32_t sqn_only;
    };
} __attribute__ ((packed)) PfcpHeader;
    ''')

    definedList = []
    ieDesTable = []
    for ieName, ieVal in ie_type_value.items():
        ieDesTable.append([ieVal, f'sizeof({ieName})', 1, 0, []]) # set default as TlvOctet struct

    table = doc29244_7_para.tables[0]
    ieTable = table
    for i, row in enumerate(table.rows):
        if (i == 0 or i == 1):
            continue
        if row.cells[0].paragraphs[0].text.isdigit():
            PFCPMessageHeaderFd.fprint('#define ' + snackCase(row.cells[1].paragraphs[0].text) + \
            ' ' + row.cells[0].paragraphs[0].text)
    PFCPMessageHeaderFd.fprint('')

    for key in ie_type_value:
        ieName, ieValue = key, ie_type_value[key]
        PFCPMessageHeaderFd.fprint(f'#define PFCP_{ieName}_TYPE {ieValue}')
    PFCPMessageHeaderFd.fprint('')

    for ieName, ieVal in ie_type_value.items():
        if ieName not in chapter7UsedIESet:
            PFCPMessageHeaderFd.fprint(f'typedef TlvOctet {ieName};')
            definedList.append(ieName)
            #ieDesTable.append([ie_type_value[ieName], f'sizeof({ieName})', 1, 0, []])
    PFCPMessageHeaderFd.fprint('')

    ieTypeNotDefinedList = []
    for tableName, table in zip(chapter7TitleList[1:], doc29244_7.tables) :
        tableName = formatString(tableName)
        ieIn = re.compile("^.*IEIn.*")
        if tableName == "UpdateBARIEInPFCPSessionReportResponse":
            tableName = "UpdateBARPFCPSessionReportResponse"
        elif tableName == "UserPlanePathFailure":
            tableName = "UserPlanePathFailureReport"
        elif tableName == "PFD":
            tableName = "PFDContext"
        elif ieIn.match(tableName):
            #print("============", tableName, tableName[:tableName.find("IEIn")])
            tableName = tableName[:tableName.find("IEIn")]
        elif tableName == 'RemoveQERIEPFCPSessionModificationRequest':
            tableName = tableName[:tableName.find("IE")]
        ieTypeNotDefined = False
        # check if exist not defined ie
        for ie in table :
            try :
                ieName = large2smallCamalCast(formatString(ie['Information elements']))
            except :
                ieName = 'NoIEName'
                print(f'[warning] No IE name in {tableName}')
            try :
                ieType = formatString(ie['IE Type'])
            except:
                print('NoIEType')
            if ieType not in definedList:
                ieTypeNotDefined = True
                break
        if ieTypeNotDefined:
            tmpTuple = [tableName, []]
            for ie in table:
                try:
                    ieName = large2smallCamalCast(formatString(ie['Information elements']))
                except:
                    print(f'No IE name in {tableName}')
                    continue
                try:
                    ieType = formatString(ie['IE Type'])
                except:
                    print('No IE type')
                try:
                    if ieNameMapping.get(ieType):
                        ieType = ieNameMapping[ieType]
                    elif ieType in specialCase:
                        ieType+= tableName
                except:
                    print("???")
                tmpTuple[1].append((ieType, ieName))
            ieTypeNotDefinedList.append(tmpTuple)
        else:
            definedList.append(tableName)
            PFCPMessageHeaderFd.fprint(f'typedef struct _{tableName} ' +'{')
            PFCPMessageHeaderFd.indent(+1)
            PFCPMessageHeaderFd.fprint('unsigned long presence;')
            ieNum = 0
            ieList = []
            for ie in table:
                try :
                    ieName = large2smallCamalCast(formatString(ie['Information elements']))
                except :
                    ieName = 'NoIEName'
                    print(f'[warning] No IE name in {tableName}')
                try :
                    ieType = formatString(ie['IE Type'])
                except :
                    ieType = 'NoIEType'
                    print(f'[warning] No IE {ieName} type in {tableName}')
                try :
                    if ieNameMapping.get(ieType) :
                        ieType = ieNameMapping[ieType]
                    elif ieType in specialCase:
                        ieType += tableName
                except :
                    print('[warning] Cannot get ieType from ieNameMapping:', ieType)
                if ie_type_value.get(ieType) != None:
                    ieList.append(ie_type_value.get(ieType))
                    ieNum += 1
                else:
                    print("IE value cannot find:", ieType, ieName, tableName)
                PFCPMessageHeaderFd.fprint(f'{ieType} {ieName};')
            if ie_type_value.get(tableName) != None:
                ieDesTable[ie_type_value.get(tableName)] = [ie_type_value.get(tableName), f'sizeof({tableName})', 0, ieNum, ieList]
            else:
                print(tableName, "not ie")
                ieDesTable.append(
                    [0, f'sizeof({tableName})', 0, ieNum, ieList])

            PFCPMessageHeaderFd.indent(-1)
            PFCPMessageHeaderFd.fprint(
                '} __attribute__((packed)) ' + f'{tableName};\n')

    for table in ieTypeNotDefinedList:
        tableName = table[0]
        ieTypeNotDefined = False
        for ie in table[1]:
            ieType = ie[0]
            if ieType not in definedList:
                ieTypeNotDefined = True
                break
        if ieTypeNotDefined:
            ieTypeNotDefinedList.append(table)
        else:
            definedList.append(tableName)
            PFCPMessageHeaderFd.fprint(f'typedef struct _{tableName} ' +'{')
            PFCPMessageHeaderFd.indent(+1)
            PFCPMessageHeaderFd.fprint('unsigned long presence;')
            ieNum = 0
            ieList = []
            for ie in table[1]:
                ieType = ie[0]
                ieName = ie[1]
                #ieNum += 1
                #ieList.append(definedList.index(ieType))
                if ie_type_value.get(ieType) != None:
                    ieList.append(ie_type_value.get(ieType))
                    ieNum += 1
                else:
                    print("IE value cannot find:", ieType)
                PFCPMessageHeaderFd.fprint(f'{ieType} {ieName};')
            if ie_type_value.get(tableName) != None:
                ieDesTable[ie_type_value.get(tableName)] = [ie_type_value.get(tableName), f'sizeof({tableName})', 0, ieNum, ieList]
                #ieDesTable.append([ie_type_value[tableName], f'sizeof({tableName})', 0, ieNum, ieList])
            else:
                print(tableName, "not ie")
                ieDesTable.append([0, f'sizeof({tableName})', 0, ieNum, ieList])
            PFCPMessageHeaderFd.indent(-1)
            PFCPMessageHeaderFd.fprint(
                '} __attribute__((packed)) ' + f'{tableName};\n')


    # PfcpMessage type
    PFCPMessageHeaderFd.fprint("")
    PFCPMessageHeaderFd.fprint("typedef struct _PfcpMessage {")
    PFCPMessageHeaderFd.indent(+1)
    PFCPMessageHeaderFd.fprint("PfcpHeader header;")
    PFCPMessageHeaderFd.fprint("union {")
    PFCPMessageHeaderFd.indent(+1)
    for i, row in enumerate(ieTable.rows):
        if (i == 0 or i == 1 or i == 2):
            continue
        if row.cells[0].paragraphs[0].text.isdigit():
            msg = snackCase(row.cells[1].paragraphs[0].text)
            if msg == "PFCP_VERSION_NOT_SUPPORTED_RESPONSE":
                continue
            elif msg[:7] == "PFCPPFD":
                PFCPMessageHeaderFd.fprint(f'{"PFCPPFD"+snack2CamalCast(msg[8:])} {large2smallCamalCast("PFCPPFD"+snack2CamalCast(msg[8:]))};')
            elif msg[5:10] == "HEART":
                PFCPMessageHeaderFd.fprint(f'{snack2CamalCast(msg[5:])} {large2smallCamalCast(snack2CamalCast(msg[5:]))};')
            else:
                PFCPMessageHeaderFd.fprint(f'{"PFCP"+snack2CamalCast(msg[5:])} {large2smallCamalCast("PFCP"+snack2CamalCast(msg[5:]))};')

    PFCPMessageHeaderFd.indent(-1)
    PFCPMessageHeaderFd.fprint("};")
    PFCPMessageHeaderFd.indent(-1)
    PFCPMessageHeaderFd.fprint("} PfcpMessage;")

    # encode & decode function declear
    PFCPMessageHeaderFd.fprint("")
    PFCPMessageHeaderFd.fprint("Status PfcpParseMessage(PfcpMessage *pfcpMessage, Bufblk *buf);")
    PFCPMessageHeaderFd.fprint("")
    PFCPMessageHeaderFd.fprint("Status PfcpBuildMessage(Bufblk **bufBlkPtr, PfcpMessage *pfcpMessage);")

    PFCPMessageHeaderFd.fprint("""
#ifdef __cplusplus
}
#endif /* __cplusplus */

#endif /* __PFCP_MESSAGE_H__ */
""")

    PFCPMessageSourceFd = FileOutput('pfcp_message.c')
    PFCPMessageSourceFd.fprint('''#define TRACE_MODULE _pfcp_message

#include <endian.h>
#include <string.h>
#include <netinet/in.h>

#include "utlt_debug.h"
#include "utlt_buff.h"

#include "pfcp_message.h"
''')

    PFCPMessageSourceFd.fprint(
        "static IeDescription ieDescriptionTable[] = {\\")
    for ieDes in ieDesTable:
        tmpStr = '{'
        idx = 0
        for idx, ie in enumerate(ieDes[4]):
            if idx != 0:
                tmpStr += ', '
            tmpStr += str(ie)
            idx += 1
        while idx < 35:
            if idx != 0:
                tmpStr += ', 0'
            else:
                tmpStr += '0'
            idx += 1
        tmpStr += '}'
        PFCPMessageSourceFd.fprint(
            '{' + f'{ieDes[0]}, {ieDes[1]}, {ieDes[2]}, {ieDes[3]}, {tmpStr}' + '}, \\')
    PFCPMessageSourceFd.fprint("};")
    PFCPMessageSourceFd.fprint('''
_Bool dbf = 0;

int _TlvParseMessage(void * msg, IeDescription * msgDes, void * buff, int buffLen) {
    int msgPivot = 0; // msg (struct) offset
    //void *root = buff;
    int buffOffset = 0; // buff offset
    int idx;
    for (idx = 0; idx < msgDes->numToParse; ++idx) {
        if (dbf) { if (ieDescriptionTable[msgDes->next[idx]].msgType == 57) {
                UTLT_Warning("Get F-SEID");
            } }
        IeDescription *ieDes = &ieDescriptionTable[msgDes->next[idx]];
        uint16_t type;
        uint16_t length;
        memcpy(&type, buff + buffOffset, sizeof(uint16_t));
        memcpy(&length, buff + buffOffset + sizeof(uint16_t), sizeof(uint16_t));
        //type = (type>>8) + ((type&0xff)<<8);
        //length = (length>>8) + ((length&0xff)<<8);
        type = ntohs(type);
        length = ntohs(length);
        if (dbf) { UTLT_Info("type: %d, len: %d", type, length); }
        if (type != ieDes->msgType) {
            if (dbf) { UTLT_Warning("%d not present, type: %d", ieDes->msgType, type); }
            // not present
            (*(unsigned long*)(msg + msgPivot)) = 0; // presence
            msgPivot += ieDes->msgLen;
            continue;
        }

        if (ieDes->isTlvObj) {
            if (dbf) { UTLT_Info("is TLV: %p", msg+msgPivot); }
            ((TlvOctet*)(msg+msgPivot))->presence = 1;
            ((TlvOctet*)(msg+msgPivot))->type = type;
            Bufblk *newBuf = BufblkAlloc(1, length);
            memcpy(newBuf->buf, buff + buffOffset + 2*sizeof(uint16_t), length);
            newBuf->len = length;
            ((TlvOctet*)(msg+msgPivot))->len = length;
            ((TlvOctet*)(msg+msgPivot))->value = newBuf->buf;
            buffOffset += sizeof(uint16_t)*2 + length;
            msgPivot += sizeof(TlvOctet);
            continue;
        } else {
            if (dbf) { UTLT_Info("not Tlv, desTB mstype: %d", ieDes->msgType); }
            // recursive
            *((unsigned long*)(msg+msgPivot)) = 1; // presence
            _TlvParseMessage(msg+msgPivot+sizeof(unsigned long), ieDes, buff + buffOffset + sizeof(uint16_t)*2, buffLen - buffOffset);
            //int size = _TlvParseMessage(msg+msgPivot, ieDes, buff + buffOffset, buffLen - buffOffset);
            buffOffset += length + sizeof(uint16_t)*2;
            msgPivot += ieDes->msgLen;
        }
    }
    return buffOffset;
}

Status PfcpParseMessage(PfcpMessage *pfcpMessage, Bufblk *bufBlk) {
    Status status = STATUS_OK;
    PfcpHeader *header = NULL;
    uint16_t size = 0;

    UTLT_Assert(pfcpMessage, return STATUS_ERROR, "Message error");
    UTLT_Assert(bufBlk, return STATUS_ERROR, "buffer error");
    UTLT_Assert(bufBlk->buf, return STATUS_ERROR, "buffer payload error");

    header = bufBlk->buf;
    UTLT_Assert(header, return STATUS_ERROR, "header hasn't get pointer");

    memset(pfcpMessage, 0, sizeof(PfcpMessage)); // clear pfcpMessage

    if (header->seidP) {
        size = PFCP_HEADER_LEN;
    } else {
        size = PFCP_HEADER_LEN - PFCP_SEID_LEN;
    }
    bufBlk->buf += size;
    bufBlk->len -= size;
    bufBlk->size -= size;

    memcpy(&pfcpMessage->header, bufBlk->buf - size, size);

    if (header->seidP) {
        pfcpMessage->header.seid = be64toh(pfcpMessage->header.seid);
    } else { // not sure what is this for
        pfcpMessage->header.sqn = pfcpMessage->header.sqn_only;
        pfcpMessage->header.sqn_only = pfcpMessage->header.sqn_only;
    }

    if (bufBlk->len == 0) {
        return STATUS_OK;
    }

    switch(pfcpMessage->header.type) {''')
    PFCPMessageSourceFd.indent(+2)
    for i, row in enumerate(ieTable.rows):
        if (i == 0 or i == 1 or i == 2):
            continue
        if row.cells[0].paragraphs[0].text.isdigit():
            msg = snackCase(row.cells[1].paragraphs[0].text)
            if msg == "PFCP_VERSION_NOT_SUPPORTED_RESPONSE":
                PFCPMessageSourceFd.fprint(f'case {msg}:')
                PFCPMessageSourceFd.indent(+1)
            elif msg[:7] == "PFCPPFD":
                PFCPMessageSourceFd.fprint(f'case {msg}:')
                PFCPMessageSourceFd.indent(+1)
                PFCPMessageSourceFd.fprint(f'pfcpMessage->{large2smallCamalCast("PFCPPFD"+snack2CamalCast(msg[8:]))}.presence = 1;')
                PFCPMessageSourceFd.fprint(f'_TlvParseMessage((unsigned long *)&pfcpMessage->{large2smallCamalCast("PFCPPFD"+snack2CamalCast(msg[8:]))} + 1, &ieDescriptionTable[{msg} + 155], bufBlk->buf, bufBlk->len);')
            elif msg[5:10] == "HEART":
                PFCPMessageSourceFd.fprint(f'case {msg}:')
                PFCPMessageSourceFd.indent(+1)
                PFCPMessageSourceFd.fprint(f'pfcpMessage->{large2smallCamalCast(snack2CamalCast(msg[5:]))}.presence = 1;')
                PFCPMessageSourceFd.fprint(f'_TlvParseMessage((unsigned long *)&pfcpMessage->{large2smallCamalCast(snack2CamalCast(msg[5:]))} + 1, &ieDescriptionTable[{msg} + 155], bufBlk->buf, bufBlk->len);')
            elif re.match("^PFCP_SESSION.*", msg) and not re.match("^PFCP_SESSION_SET.*", msg):
                PFCPMessageSourceFd.fprint(f'case {msg}:')
                PFCPMessageSourceFd.indent(+1)
                PFCPMessageSourceFd.fprint(f'pfcpMessage->{large2smallCamalCast("PFCP"+snack2CamalCast(msg[5:]))}.presence = 1;')
                PFCPMessageSourceFd.fprint(f'_TlvParseMessage((unsigned long *)&pfcpMessage->{large2smallCamalCast("PFCP"+snack2CamalCast(msg[5:]))} + 1, &ieDescriptionTable[{msg} + 155 - (50-15) - 1], bufBlk->buf, bufBlk->len);')
            else:
                PFCPMessageSourceFd.fprint(f'case {msg}:')
                PFCPMessageSourceFd.indent(+1)
                PFCPMessageSourceFd.fprint(f'pfcpMessage->{large2smallCamalCast("PFCP"+snack2CamalCast(msg[5:]))}.presence = 1;')
                if i > 13:
                    PFCPMessageSourceFd.fprint(f'_TlvParseMessage((unsigned long *)&pfcpMessage->{large2smallCamalCast("PFCP"+snack2CamalCast(msg[5:]))} + 1, &ieDescriptionTable[{msg} + 155 - 1], bufBlk->buf, bufBlk->len);')
                else:
                    PFCPMessageSourceFd.fprint(f'_TlvParseMessage((unsigned long *)&pfcpMessage->{large2smallCamalCast("PFCP"+snack2CamalCast(msg[5:]))} + 1, &ieDescriptionTable[{msg}+155], bufBlk->buf, bufBlk->len);')
            PFCPMessageSourceFd.fprint('break;')
            PFCPMessageSourceFd.indent(-1)
    PFCPMessageSourceFd.indent(-2)

    PFCPMessageSourceFd.fprint('''        default:
            UTLT_Warning("Not implmented(type:%d)", &pfcpMessage->header.type);
    }

    return status;
}

int _TlvBuildMessage(Bufblk **bufBlkPtr, void *msg, IeDescription *ieDescription) {
    //UTLT_Warning("Addr : %p", msg);
    UTLT_Assert(bufBlkPtr, return 0, "buffer error");
    UTLT_Assert(msg, return 0, "message error");
    if (*(unsigned long *)msg == 0) {
        // present bit
        //UTLT_Warning("no ie");
        return 0;
    }

    if (ieDescription->isTlvObj) {
        //UTLT_Info("TLV: type: %d, len: %d", ((TlvOctet *)msg)->type, ((TlvOctet *)msg)->len);
        //UTLT_Info("msgType: %d, msgLen: %d", ieDescription->msgType, ((TlvOctet *)msg)->len);
        int buffLen = sizeof(uint16_t) * 2 + ((TlvOctet *)msg)->len;
        *bufBlkPtr = BufblkAlloc(1, buffLen);
        uint16_t *tagPtr = (uint16_t *) ((*bufBlkPtr)->buf);
        uint16_t *lenPtr = &tagPtr[1];

        (*bufBlkPtr)->len = buffLen;
        *tagPtr = htons(ieDescription->msgType);
        *lenPtr = htons(buffLen - sizeof(uint16_t) * 2);
        memcpy((void *) &tagPtr[2], ((TlvOctet *)msg)->value, ((TlvOctet *)msg)->len);
    } else {
        UTLT_Info("not TLV");
        size_t idx;
        int msgPivot = 0;
        *bufBlkPtr = BufblkAlloc(1, sizeof(uint16_t) * 2);
        uint16_t *tagPtr = (*bufBlkPtr)->buf;
        uint16_t *lenPtr = &tagPtr[1];
        (*bufBlkPtr)->len = sizeof(uint16_t) * 2;

        *tagPtr = htons(ieDescription->msgType);
        UTLT_Warning("Check addr: tag: %p, buf: %p", tagPtr, (*bufBlkPtr)->buf);
        UTLT_Info("msgType: %u, tagPtr value: %u, first type: %u", ieDescription->msgType, ((uint16_t*)tagPtr)[0],ntohs(((uint16_t*)(*bufBlkPtr)->buf)[0]));
        *lenPtr = htons(0);

        int bufOffset = 0;
        void *msgNoPresentPtr = &((unsigned long*)msg)[1];
        for (idx = 0; idx < ieDescription->numToParse; ++idx) {
            Bufblk *tmpBufBlkPtr = NULL;
            bufOffset += _TlvBuildMessage(&tmpBufBlkPtr, &((uint8_t *)msgNoPresentPtr)[msgPivot], &ieDescriptionTable[ieDescription->next[idx]]);
            if (tmpBufBlkPtr == NULL) {
                msgPivot += ieDescriptionTable[ieDescription->next[idx]].msgLen;
                //UTLT_Info("TL type[%d], pivot %d", ieDescriptionTable[ieDescription->next[idx]].msgType, msgPivot);
                continue;
            }
            UTLT_Info("tmpBuf T: %u, L: %d", ntohs(((uint16_t *)tmpBufBlkPtr->buf)[0]), ntohs(((uint16_t *)tmpBufBlkPtr->buf)[1]));
            BufblkBuf(*bufBlkPtr, tmpBufBlkPtr);
            //UTLT_Warning("bufBlk len %d", (*bufBlkPtr)->buf);
            BufblkFree(tmpBufBlkPtr);
            msgPivot += ieDescriptionTable[ieDescription->next[idx]].msgLen;
            UTLT_Info("buff offset: %d, buff Len: %d", bufOffset, (*bufBlkPtr)->len);
        }

        *lenPtr = htons(bufOffset);
    }

    //UTLT_Warning("buf len: %d, first type: %d", (*bufBlkPtr)->len, ((uint16_t*)(*bufBlkPtr)->buf)[0]);
    return (*bufBlkPtr)->len;
}

void _PfcpBuildBody(Bufblk **bufBlkPtr, void *msg, IeDescription *ieDescription) {
    UTLT_Assert(bufBlkPtr, return, "buffer error");
    UTLT_Assert(msg, return, "message error");

    int idx;
    void *root = msg + sizeof(unsigned long);
    (*bufBlkPtr) = BufblkAlloc(1, 0);
    for (idx = 0; idx < ieDescription->numToParse; ++idx) {
        Bufblk *tmpBufBlkPtr;
        int rt = _TlvBuildMessage(&tmpBufBlkPtr, root, &ieDescriptionTable[ieDescription->next[idx]]);
        if (rt == 0) {
            root += ieDescriptionTable[ieDescription->next[idx]].msgLen;
            continue;
        }
        BufblkBuf(*bufBlkPtr, tmpBufBlkPtr);
        BufblkFree(tmpBufBlkPtr);
        root += ieDescriptionTable[ieDescription->next[idx]].msgLen;
    }
}

Status PfcpBuildMessage(Bufblk **bufBlkPtr, PfcpMessage *pfcpMessage) {
    Status status = STATUS_OK;
    UTLT_Assert(pfcpMessage, return STATUS_ERROR, "pfcpMessage error");

    switch(pfcpMessage->header.type) {''')
    
    PFCPMessageSourceFd.indent(+2)
    for i, row in enumerate(ieTable.rows):
        if (i == 0 or i == 1 or i == 2):
            continue
        if row.cells[0].paragraphs[0].text.isdigit():
            msg = snackCase(row.cells[1].paragraphs[0].text)
            if msg == "PFCP_VERSION_NOT_SUPPORTED_RESPONSE":
                PFCPMessageSourceFd.fprint(f'case {msg}:')
                PFCPMessageSourceFd.indent(+1)
            elif msg[:7] == "PFCPPFD":
                PFCPMessageSourceFd.fprint(f'case {msg}:')
                PFCPMessageSourceFd.indent(+1)
                PFCPMessageSourceFd.fprint(f'_PfcpBuildBody(bufBlkPtr, &pfcpMessage->{large2smallCamalCast("PFCPPFD"+snack2CamalCast(msg[8:]))}, &ieDescriptionTable[{msg} + 155]);')
            elif msg[5:10] == "HEART":
                PFCPMessageSourceFd.fprint(f'case {msg}:')
                PFCPMessageSourceFd.indent(+1)
                PFCPMessageSourceFd.fprint(f'_PfcpBuildBody(bufBlkPtr, &pfcpMessage->{large2smallCamalCast(snack2CamalCast(msg[5:]))}, &ieDescriptionTable[{msg} + 155]);')
            elif re.match("^PFCP_SESSION.*", msg) and not re.match("^PFCP_SESSION_SET.*", msg):
                PFCPMessageSourceFd.fprint(f'case {msg}:')
                PFCPMessageSourceFd.indent(+1)
                PFCPMessageSourceFd.fprint(f'_PfcpBuildBody(bufBlkPtr, &pfcpMessage->{large2smallCamalCast("PFCP"+snack2CamalCast(msg[5:]))}, &ieDescriptionTable[{msg} + 155 - (50-15) - 1]);')
            else:
                PFCPMessageSourceFd.fprint(f'case {msg}:')
                PFCPMessageSourceFd.indent(+1)
                if i > 13:
                    PFCPMessageSourceFd.fprint(f'_PfcpBuildBody(bufBlkPtr, &pfcpMessage->{large2smallCamalCast("PFCP"+snack2CamalCast(msg[5:]))}, &ieDescriptionTable[{msg} + 155 - 1]);')
                else:
                    PFCPMessageSourceFd.fprint(f'_PfcpBuildBody(bufBlkPtr, &pfcpMessage->{large2smallCamalCast("PFCP"+snack2CamalCast(msg[5:]))}, &ieDescriptionTable[{msg} + 155]);')
            PFCPMessageSourceFd.fprint('break;')
            PFCPMessageSourceFd.indent(-1)
    PFCPMessageSourceFd.indent(-2)

    PFCPMessageSourceFd.fprint('''        default:
            UTLT_Warning("Not implmented(type:%d)", &pfcpMessage->header.type);
    }
    return status;
}
    ''')

